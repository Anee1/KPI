import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def generer_document_word(donnees):
    """Génère le fichier Word à partir des données du formulaire."""
    doc = Document()
    
    # Titre principal
    titre = doc.add_heading('COMPTE RENDU DE VISITE COMMERCIALE', 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Informations Générales
    doc.add_heading('1. Informations Générales', level=1)
    p = doc.add_paragraph()
    p.add_run('Date de visite : ').bold = True
    p.add_run(f"{donnees['Date']}\t\t")
    p.add_run('Heure : ').bold = True
    p.add_run(f"{donnees['Heure_debut']} à {donnees['Heure_fin']}\n")
    
    p.add_run('Commercial : ').bold = True
    p.add_run(f"{donnees['Commercial']}\t\t")
    p.add_run('Fonction : ').bold = True
    p.add_run(f"{donnees['Fonction_Comm']}\n")
    
    p.add_run('Interlocuteur : ').bold = True
    p.add_run(f"{donnees['Interlocuteur']}\t\t")
    p.add_run('Fonction : ').bold = True
    p.add_run(f"{donnees['Fonction_Inter']}\n")
    
    p.add_run('Nom du Prospect : ').bold = True
    p.add_run(f"{donnees['Prospect']}\n")
    p.add_run('Lieu de la visite : ').bold = True
    p.add_run(f"{donnees['Lieu']}")

    # 2. Résumé de l'échange
    doc.add_heading('2. Synthèse de la Visite', level=1)
    doc.add_paragraph('Objectif de la visite :').runs[0].bold = True
    doc.add_paragraph(donnees['Objectif'])
    
    doc.add_paragraph('Résumé de l’échange :').runs[0].bold = True
    doc.add_paragraph(donnees['Resume'])

    # 3. Informations clés
    doc.add_heading('3. Informations Clés sur le Client', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = 'Budget d’investissement :'
    cells[1].text = f"{donnees['Budget']:,.0f} FCFA"
    
    cells = table.rows[1].cells
    cells[0].text = 'Échéance prévue :'
    cells[1].text = str(donnees['Echeance'])
    
    cells = table.rows[2].cells
    cells[0].text = 'Décisionnaires identifiés :'
    cells[1].text = donnees['Decisionnaire']
    
    cells = table.rows[3].cells
    cells[0].text = 'Niveau d’intérêt :'
    cells[1].text = donnees['Interet']

    # 4. Spécificités
    doc.add_heading('4. Spécificités et Attentes', level=1)
    doc.add_paragraph('Contraintes spécifiques :').runs[0].bold = True
    doc.add_paragraph(donnees['Contraintes'])
    
    doc.add_paragraph('Attentes (impact, innovation) :').runs[0].bold = True
    doc.add_paragraph(donnees['Attentes'])
    
    doc.add_paragraph('Proposition / Piste de collaboration :').runs[0].bold = True
    doc.add_paragraph(donnees['Proposition'])

    # 5. Plan d'action
    doc.add_heading('5. Prochaines Actions et Livrables', level=1)
    doc.add_paragraph('Actions à mener :').runs[0].bold = True
    doc.add_paragraph(donnees['Actions'])
    
    doc.add_paragraph('Livrables attendus :').runs[0].bold = True
    doc.add_paragraph(donnees['Livrables'])

    # Sauvegarde en mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def afficher_formulaire_crv():
    st.title("📝 Saisie du Compte Rendu de Visite (CRV)")
    
    # --- DEBUT DU FORMULAIRE ---
    with st.form("form_crv"):
        st.subheader("1. Informations de la visite")
        col1, col2, col3 = st.columns(3)
        date_visite = col1.date_input("Date de visite")
        heure_debut = col2.time_input("Heure de début")
        heure_fin = col3.time_input("Heure de fin")
        
        col4, col5 = st.columns(2)
        commercial = col4.selectbox("Commercial", ["Grace", "Nelly", "Imamiah"])
        fonction_comm = col5.text_input("Fonction (Commercial)", value="Analyste / Chargé d'Affaires")
        
        col6, col7 = st.columns(2)
        interlocuteur = col6.text_input("Nom de l'interlocuteur")
        fonction_inter = col7.text_input("Fonction de l'interlocuteur")
        
        nom_prospect = st.text_input("Nom du Prospect / Entreprise")
        lieu_visite = st.text_input("Lieu de la visite")

        st.subheader("2. Détails de l'échange")
        objectif = st.text_input("Objectif de la visite")
        resume = st.text_area("Résumé de l’échange", height=150)

        st.subheader("3. Informations clés")
        col8, col9 = st.columns(2)
        budget = col8.number_input("Budget d’investissement (FCFA)", min_value=0)
        echeance = col9.date_input("Échéance prévue")
        
        decisionnaire = st.text_input("Décisionnaires identifiés")
        interet = st.radio("Niveau d’intérêt à nos produits", ["Faible", "Moyen", "Élevé", "RAS"], horizontal=True)

        st.subheader("4. Spécificités")
        contraintes = st.text_area("Contraintes spécifiques (budgétaires, délais, procédures)")
        attentes = st.text_area("Attentes en termes d’impact, d’innovation")
        proposition = st.text_area("Proposition ou piste de collaboration")

        st.subheader("5. Prochaines actions à mener")
        actions = st.text_area("Actions à mener (Ex: Envoyer la plaquette, relancer le client le XX/XX)")
        livrables = st.text_area("Documents ou livrables attendus (dossier, note d’information...)")

        # Le seul bouton autorisé dans le formulaire
        soumis = st.form_submit_button("✅ Générer le Compte Rendu (Word)")
    # --- FIN DU FORMULAIRE (Attention à l'indentation ci-dessous) ---


    # --- LOGIQUE HORS DU FORMULAIRE ---
    if soumis:
        # Regroupement des données
        donnees_crv = {
            "Date": date_visite.strftime("%d/%m/%Y"),
            "Heure_debut": heure_debut.strftime("%H:%M"),
            "Heure_fin": heure_fin.strftime("%H:%M"),
            "Commercial": commercial,
            "Fonction_Comm": fonction_comm,
            "Interlocuteur": interlocuteur,
            "Fonction_Inter": fonction_inter,
            "Prospect": nom_prospect,
            "Lieu": lieu_visite,
            "Objectif": objectif,
            "Resume": resume,
            "Budget": budget,
            "Echeance": echeance.strftime("%d/%m/%Y"),
            "Decisionnaire": decisionnaire,
            "Interet": interet,
            "Contraintes": contraintes,
            "Attentes": attentes,
            "Proposition": proposition,
            "Actions": actions,
            "Livrables": livrables
        }
        
        # Génération du document en mémoire
        fichier_word = generer_document_word(donnees_crv)
        nom_fichier = f"CRV_{nom_prospect.replace(' ', '_')}_{date_visite.strftime('%Y%m%d')}.docx"
        
        # On sauvegarde le document dans la mémoire tampon (session_state) de Streamlit
        st.session_state['doc_crv'] = fichier_word
        st.session_state['nom_fichier_crv'] = nom_fichier
        
        st.success("Le compte rendu a été généré avec succès !")

    # Si le document existe en mémoire, on affiche le bouton de téléchargement
    if 'doc_crv' in st.session_state:
        st.download_button(
            label="📥 Télécharger le CRV (Format Word)",
            data=st.session_state['doc_crv'],
            file_name=st.session_state['nom_fichier_crv'],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
